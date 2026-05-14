"""
StyleManager 
"""

import hashlib
from docx import Document
from docx.shared import Inches

from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement, parse_xml
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re
import logging
from copy import deepcopy

# Configure logging
logger = logging.getLogger(__name__)


def get_image_hash(blob):
    """Generate MD5 hash for a binary blob (image)."""
    return hashlib.md5(blob).hexdigest()


# Relationship namespace
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'


class StyleManager:
    def __init__(self, template_path: str, font_family: str = 'Calibri', font_size: int = 11):
        self.template_path = template_path
        self.template = Document(template_path)
        
        # User-selected settings
        self.font_family = font_family
        self.BODY_FONT = font_family 
        self.font_size = font_size
        
        self.BODY_COLOR = "000000"
        self.COLOR_H1 = "0F1F38" # Deep Navy (Unique Start)
        self.COLOR_H2 = "183A60" # Classic Navy (Previous Primary)
        self.COLOR_H3 = "82ADCF" # Soft Blue
        self.COLOR_H4 = "374151" # Slate Gray (Smallest)
        self.TABLE_BG_LIGHT = "EDF4FC"
        
        # Calculate sizes based on user's choice (in half-points)
        self.BODY_SIZE = font_size * 2              # 11pt → 22
        self.MIN_SIZE = self.BODY_SIZE
        self.H1_SIZE = int(font_size * 1.45 * 2)    # 11pt → 16pt (32 half-pts)
        self.H2_SIZE = int(font_size * 1.27 * 2)    # 11pt → 14pt (28 half-pts)
        self.H3_SIZE = int(font_size * 1.18 * 2)    # 11pt → 13pt (26 half-pts)
        self.H4_SIZE = int(font_size * 1.09 * 2)

    def apply_list_styling(self, paragraph, level: int = 0):
        """
        Normalizes list indentation and aligns bullet markers cleanly.
        """
        pf = paragraph.paragraph_format
        style_name = paragraph.style.name if paragraph.style else ''
        is_bullet_list = 'Bullet' in style_name
        level = max(0, int(level or 0))
        
        # Reset inherited tab stops so marker/text spacing is predictable.
        pPr = paragraph._element.get_or_add_pPr()
        for old_tabs in pPr.findall(qn('w:tabs')):
            pPr.remove(old_tabs)
        
        # Keep a compact hanging indent while preserving nesting depth.
        if is_bullet_list:
            # Nudge bullets and text slightly to the right for a cleaner visual margin.
            left_indent = 0.32 + (level * 0.18)
            pf.left_indent = Inches(left_indent)
            pf.first_line_indent = Inches(-0.18)
            pf.tab_stops.add_tab_stop(Inches(left_indent))
            self._set_list_marker_alignment(paragraph, marker_alignment='right')
        else:
            left_indent = 0.20 + (level * 0.22)
            pf.left_indent = Inches(left_indent)
            pf.first_line_indent = Inches(-0.20)
            pf.tab_stops.add_tab_stop(Inches(left_indent))
            self._set_list_marker_alignment(paragraph, marker_alignment='left')
        
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Force font for the list marker (bullet/number)
        pPr = paragraph._element.get_or_add_pPr()
        rPr = pPr.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            pPr.append(rPr)
        for old_fonts in rPr.findall(qn('w:rFonts')):
            rPr.remove(old_fonts)
        fonts = OxmlElement('w:rFonts')
        fonts.set(qn('w:ascii'), self.font_family)
        fonts.set(qn('w:hAnsi'), self.font_family)
        fonts.set(qn('w:cs'), self.font_family)
        rPr.append(fonts)

    def _set_list_marker_alignment(self, paragraph, marker_alignment: str = 'left'):
        """Update the numbering definition so Word aligns the list marker as requested."""
        pPr = paragraph._element.find(qn('w:pPr'))
        if pPr is None:
            return

        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            return

        num_id_el = numPr.find(qn('w:numId'))
        ilvl_el = numPr.find(qn('w:ilvl'))
        if num_id_el is None:
            return

        try:
            num_id = int(num_id_el.get(qn('w:val')))
            ilvl = int(ilvl_el.get(qn('w:val'), '0')) if ilvl_el is not None else 0
            numbering = paragraph.part.numbering_part.numbering_definitions._numbering
            num = numbering.num_having_numId(num_id)
            abstract_num_id = int(num.abstractNumId.val)

            abstract_num = None
            for candidate in numbering.findall(qn('w:abstractNum')):
                if candidate.get(qn('w:abstractNumId')) == str(abstract_num_id):
                    abstract_num = candidate
                    break
            if abstract_num is None:
                return

            lvl = None
            for candidate in abstract_num.findall(qn('w:lvl')):
                if candidate.get(qn('w:ilvl')) == str(ilvl):
                    lvl = candidate
                    break
            if lvl is None:
                return

            lvl_jc = lvl.find(qn('w:lvlJc'))
            if lvl_jc is None:
                lvl_jc = OxmlElement('w:lvlJc')
                lvl.append(lvl_jc)
            lvl_jc.set(qn('w:val'), marker_alignment)
        except Exception as exc:
            logger.debug(f'Unable to set list marker alignment: {exc}')

    def apply_template_styles(self, doc: Document) -> Document:
        """Single entry point — full branding pipeline."""
        logger.info('Applying branding...')
        self._remove_document_protection(doc)
        self._remove_leading_empty_paragraphs(doc)
        self._override_style_sizes(doc)
        self._brand_body_paragraphs(doc)
        self._add_h1_borders(doc)
        self._add_section_separators(doc)
        self._style_tables(doc)
        self._apply_page_layout_rules(doc)
        self._apply_heading_indents(doc)
        self._copy_header_footer(doc)
        self._add_page_number(doc)
        self._adjust_page_margins(doc)

       

        logger.info('[StyleManager] ✓ Done')
        return doc

    
    # Style-level sizes (styles.xml, not run-by-run)
    

    def _override_style_sizes(self, doc: Document):
        config = {
            'Heading1':     {'sz': self.H1_SIZE,   'bold': True,  'color': self.COLOR_H1},
            'Heading 1':    {'sz': self.H1_SIZE,   'bold': True,  'color': self.COLOR_H1},
            'Heading2':     {'sz': self.H2_SIZE,   'bold': True,  'color': self.COLOR_H2},
            'Heading 2':    {'sz': self.H2_SIZE,   'bold': True,  'color': self.COLOR_H2},
            'Heading3':     {'sz': self.H3_SIZE,   'bold': True,  'color': self.COLOR_H3},
            'Heading 3':    {'sz': self.H3_SIZE,   'bold': True,  'color': self.COLOR_H3},
            # Linked CHARACTER styles
            'Heading1Char': {'sz': self.H1_SIZE,   'bold': True,  'color': self.COLOR_H1},
            'Heading2Char': {'sz': self.H2_SIZE,   'bold': True,  'color': self.COLOR_H2},
            'Heading3Char': {'sz': self.H3_SIZE,   'bold': True,  'color': self.COLOR_H3},
            'Heading4':     {'sz': self.H4_SIZE,   'bold': True,  'color': self.COLOR_H4},
            'Heading 4':    {'sz': self.H4_SIZE,   'bold': True,  'color': self.COLOR_H4},
            'Heading4Char': {'sz': self.H4_SIZE,   'bold': True,  'color': self.COLOR_H4},
        }

        styles_elem = doc.part.styles._element
        for style_el in styles_elem.findall(qn('w:style')):
            style_id   = style_el.get(qn('w:styleId'), '')
            name_el    = style_el.find(qn('w:name'))
            style_name = name_el.get(qn('w:val'), '') if name_el is not None else ''

            cfg = config.get(style_id) or config.get(style_name)
            if cfg is None:
                continue

            rPr = style_el.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                style_el.append(rPr)

            for t in ('w:sz', 'w:szCs'):
                for e in rPr.findall(qn(t)):
                    rPr.remove(e)
            sz = OxmlElement('w:sz');   sz.set(qn('w:val'), str(cfg['sz']));   rPr.append(sz)
            szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), str(cfg['sz'])); rPr.append(szCs)

            for e in rPr.findall(qn('w:color')): rPr.remove(e)
            c = OxmlElement('w:color'); c.set(qn('w:val'), cfg.get('color', '000000')); rPr.append(c)

            for e in rPr.findall(qn('w:rFonts')): rPr.remove(e)
            f = OxmlElement('w:rFonts')
            f.set(qn('w:ascii'), self.font_family); f.set(qn('w:hAnsi'), self.font_family)
            f.set(qn('w:cs'), self.font_family);   rPr.append(f)

            for e in rPr.findall(qn('w:b')): rPr.remove(e)
            if cfg['bold']:
                rPr.append(OxmlElement('w:b'))

            logger.debug(f'  Style "{style_id or style_name}" → '
                f'{cfg["sz"] / 2}pt {self.font_family} bold={cfg["bold"]}')
            
            pPr = style_el.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                style_el.append(pPr)

            for e in pPr.findall(qn('w:spacing')):
                pPr.remove(e)

            spacing = OxmlElement('w:spacing')

            if cfg['bold']:
                spacing.set(qn('w:before'), '120')
                spacing.set(qn('w:after'), '120')
                spacing.set(qn('w:line'), '240')
                spacing.set(qn('w:lineRule'), 'auto')
            else:
                spacing.set(qn('w:before'), '0')
                spacing.set(qn('w:after'), '20')
                spacing.set(qn('w:line'), '220')
                spacing.set(qn('w:lineRule'), 'auto')

            pPr.append(spacing)

            if cfg['bold']:   
                for e in pPr.findall(qn('w:keepNext')):
                    pPr.remove(e)
                pPr.append(OxmlElement('w:keepNext'))
        
        # Force global default font in styles (Normal style and docDefaults)
        self._set_global_font_defaults(doc)

    def _set_global_font_defaults(self, doc: Document):
        """Overrides the document-wide default font in styles.xml."""
        styles_elem = doc.part.styles._element
        
        # 1. Update docDefaults
        doc_defaults = styles_elem.find(qn('w:docDefaults'))
        if doc_defaults is not None:
            rPrDefault = doc_defaults.find(qn('w:rPrDefault'))
            if rPrDefault is not None:
                rPr = rPrDefault.find(qn('w:rPr'))
                if rPr is not None:
                    for old_fonts in rPr.findall(qn('w:rFonts')):
                        rPr.remove(old_fonts)
                    f = OxmlElement('w:rFonts')
                    f.set(qn('w:ascii'), self.font_family)
                    f.set(qn('w:hAnsi'), self.font_family)
                    f.set(qn('w:cs'), self.font_family)
                    rPr.append(f)

        # 2. Update 'Normal' style specifically
        for style_el in styles_elem.findall(qn('w:style')):
            style_id = style_el.get(qn('w:styleId'), '')
            if style_id == 'Normal':
                rPr = style_el.find(qn('w:rPr'))
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    style_el.append(rPr)
                for old_fonts in rPr.findall(qn('w:rFonts')):
                    rPr.remove(old_fonts)
                f = OxmlElement('w:rFonts')
                f.set(qn('w:ascii'), self.font_family)
                f.set(qn('w:hAnsi'), self.font_family)
                f.set(qn('w:cs'), self.font_family)
                rPr.append(f)

    def _apply_heading_indents(self, doc):
        """Apply indentation to headings, body text, and tables using fixed baselines"""
        current_baseline = 0
        
        # Iterate through the document body in order to maintain context (e.g. for tables)
        for element in doc.element.body:
            if isinstance(element, CT_P):
                para = Paragraph(element, doc)
                style = para.style.name if para.style else ''
                pPr = para._p.get_or_add_pPr()
                
                # Remove old indent
                for e in pPr.findall(qn('w:ind')):
                    pPr.remove(e)
                
                # 1. Update baseline based on headings (Removed cascading indents per user request)
                indent_left = 0
                current_baseline = 0
                
                is_separator = False
                pBdr = pPr.find(qn('w:pBdr'))
                if pBdr is not None and pBdr.find(qn('w:bottom')) is not None and not para.text.strip() and not style.startswith('Heading'):
                    is_separator = True
                    indent_left = 0
                
                # 2. Handle Lists (Detect by Style name OR explicit numPr)
                numPr = pPr.find(qn('w:numPr'))
                if 'List' in style or numPr is not None:
                    self.apply_list_styling(para)
                else:
                    ind = OxmlElement('w:ind')
                    ind.set(qn('w:left'), str(indent_left))
                    if is_separator:
                        ind.set(qn('w:right'), '0')
                    pPr.append(ind)
                
                # 🟢 KEEP CAPTIONS WITH TABLES (Page break prevention)
                if para.text.strip().startswith("Table"):
                    para.paragraph_format.keep_with_next = True

            elif isinstance(element, CT_Tbl):
                table = Table(element, doc)
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                
                # 🟢 Match current_baseline (indentation of the section it belongs to)
                tblPr = table._element.tblPr
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    table._element.insert(0, tblPr)
                
                tblInd = OxmlElement('w:tblInd')
                tblInd.set(qn('w:w'), str(current_baseline))
                tblInd.set(qn('w:type'), 'dxa')
                
                # Clear existing indents
                for old_ind in tblPr.findall(qn('w:tblInd')):
                    tblPr.remove(old_ind)
                tblPr.append(tblInd)

    def _style_tables(self, doc: Document):
        """Apply colors to tables: header #1A3C5E/White, zebra rows #EDF4FC/#FFFFFF, borders #D2DDED."""
        logger.info('Styling tables...')
        border_color = "D2DDED"
        header_bg = "1A3C5E"
        
        for table in doc.tables:
            # 1. Set Table Borders
            tblPr = table._element.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                table._element.insert(0, tblPr)

            tblBorders = OxmlElement('w:tblBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'single')
                b.set(qn('w:sz'), '4') # 0.5 pt
                b.set(qn('w:space'), '0')
                b.set(qn('w:color'), border_color)
                tblBorders.append(b)
            
            # Remove old borders
            for old_b in tblPr.findall(qn('w:tblBorders')):
                tblPr.remove(old_b)
            tblPr.append(tblBorders)

            for i, row in enumerate(table.rows):
                # PREVENT ROW SPLITTING (Page break prevention)
                row.allow_break_across_pages = False
                
                is_header = (i == 0)
                bg_color = header_bg if is_header else (self.TABLE_BG_LIGHT if i % 2 != 0 else "FFFFFF")
                
                for cell in row.cells:
                    # Set background color
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), bg_color)
                    shd.set(qn('w:val'), 'clear')
                    for old_shd in tcPr.findall(qn('w:shd')):
                        tcPr.remove(old_shd)
                    tcPr.append(shd)
                    
                    # Vertical Alignment for header
                    if is_header:
                        v_align = OxmlElement('w:vAlign')
                        v_align.set(qn('w:val'), 'center')
                        for old_v in tcPr.findall(qn('w:vAlign')):
                            tcPr.remove(old_v)
                        tcPr.append(v_align)
                    
                    # Set text color and formatting
                    for para in cell.paragraphs:
                        if is_header:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            # KEEP HEADER WITH DATA (Page break prevention)
                            para.paragraph_format.keep_with_next = True
                        
                        for run in para.runs:
                            rPr = run._r.get_or_add_rPr()
                            for old_color in rPr.findall(qn('w:color')):
                                rPr.remove(old_color)
                            color_el = OxmlElement('w:color')
                            color_el.set(qn('w:val'), 'FFFFFF' if is_header else self.BODY_COLOR)
                            rPr.append(color_el)
                            
                            if is_header:
                                # Ensure header is bold
                                for old_b in rPr.findall(qn('w:b')):
                                    rPr.remove(old_b)
                                rPr.append(OxmlElement('w:b'))

    def _add_h1_borders(self, doc: Document):
        """Add bottom border only to every H1."""
        logger.info('Adding H1 borders...')
        
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ''
            if style_name in ('Heading 1', 'Heading1'):
                pPr = para._p.get_or_add_pPr()
                
                # Check for existing pBdr or create new
                pBdr = pPr.find(qn('w:pBdr'))
                if pBdr is None:
                    pBdr = OxmlElement('w:pBdr')
                    pPr.append(pBdr)
                
                # Clear all existing top/bottom borders first to ensure clean state
                for side in ('top', 'bottom'):
                    for old_side in pBdr.findall(qn(f'w:{side}')):
                        pBdr.remove(old_side)
                
                # Add new bottom border only
                border = OxmlElement('w:bottom')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '6')      # 0.75 pt
                border.set(qn('w:space'), '4')   # Gap between text and line
                border.set(qn('w:color'), self.COLOR_H1)
                pBdr.append(border)

    def _add_section_separators(self, doc: Document):
        """Add a separator at the end of an H1 section (i.e., just before the next H1)."""
        logger.info('Adding section separators...')
        
        target_paras = []
        first_h1_found = False
        for i in range(len(doc.paragraphs)):
            p = doc.paragraphs[i]
            style_name = p.style.name if p.style else ''
            if style_name in ('Heading 1', 'Heading1'):
                if not first_h1_found:
                    first_h1_found = True
                    continue
                
                # Check if there is actual content before adding
                prev_p = doc.paragraphs[i-1] if i > 0 else None
                if prev_p and prev_p.text.strip() != "":
                    target_paras.append(p)
                        
        for p in target_paras:
            sep_para = p.insert_paragraph_before("")
            pPr = sep_para._p.get_or_add_pPr()
            
            pBdr = OxmlElement('w:pBdr')
            pPr.append(pBdr)
            
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6') # 0.75 pt
            bottom.set(qn('w:space'), '12')
            bottom.set(qn('w:color'), 'D9D9D9') # Subtle light gray separator
            pBdr.append(bottom)
            
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '0')
            pPr.append(spacing)

    def _copy_header_footer(self, doc: Document):
        template_section = self.template.sections[0]
        
        # Reload template fresh each time to avoid stale relationships
        self.template = Document(self.template_path)
        template_section = self.template.sections[0]
        for section in doc.sections:
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
            self._deep_copy_hdr_ftr(template_section.header, section.header, is_footer=False)
            self._deep_copy_hdr_ftr(template_section.footer, section.footer, is_footer=True)
        self._fix_duplicate_media(doc)
            
    def _add_page_number(self, doc: Document):
        
        for section in doc.sections:
            footer = section.footer

            p = footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            run = p.add_run()
            run.font.name = self.font_family

            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar1)

            instrText = OxmlElement('w:instrText')
            instrText.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            instrText.text = " PAGE "
            run._r.append(instrText)

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar2)

    def _deep_copy_hdr_ftr(self, source, target, is_footer: bool = False):
        rId_map = {}
        for old_rId, rel in source.part.rels.items():
            try:
                if 'image' in rel.reltype:
                    new_rId = target.part.relate_to(rel.target_part, rel.reltype)
                    rId_map[old_rId] = new_rId
                elif 'hyperlink' in rel.reltype:
                    new_rId = target.part.relate_to(
                        rel.target_ref, rel.reltype, is_external=True)
                    rId_map[old_rId] = new_rId
            except Exception as e:
                logger.warning(f'Rel copy warning ({old_rId}): {e}')

        target_elem = target._element
        for p in list(target_elem.findall(qn('w:p'))):
            target_elem.remove(p)

        for src_p in source._element.findall(qn('w:p')):
            # ENFORCE GAP FOR BORDERS (Blue lines overlapping fix)
            # Find pBdr (paragraph border) in src_p properties
            pPr = src_p.find(qn('w:pPr'))
            if pPr is not None:
                pBdr = pPr.find(qn('w:pBdr'))
                if pBdr is not None:
                    # Apply a gap (w:space) of 12 points to top and bottom borders
                    for border_tag in ('w:top', 'w:bottom'):
                        border = pBdr.find(qn(border_tag))
                        if border is not None:
                            border.set(qn('w:space'), '30') # Increased from 24

            # SHIFT VML SHAPES (Blue lines fix for absolute positions)
            xml_str = src_p.xml
            if '<v:rect' in xml_str:
                # Find margin-top or top in the style attribute
                def shift_vml(match):
                    style = match.group(1)
                    # Try to find margin-top:XXpt or top:XXpt
                    mt_match = re.search(r'(margin-top|top):\s*([-+]?\d*\.?\d+)(pt|in|cm|px)', style)
                    if mt_match:
                        prop = mt_match.group(1)
                        val = float(mt_match.group(2))
                        unit = mt_match.group(3)
                        new_val = val + 15 # Shift down by 15 units (usually points)
                        new_style = style.replace(mt_match.group(0), f"{prop}:{new_val}{unit}")
                        return f'style="{new_style}"'
                    return match.group(0)

                xml_str = re.sub(r'style="([^"]+)"', shift_vml, xml_str)
                new_p = parse_xml(xml_str)
            else:
                new_p = deepcopy(src_p)

            # Extra breathing room for the paragraph itself
            pPr = new_p.find(qn('w:pPr'))
            if pPr is not None:
                spacing = pPr.find(qn('w:spacing'))
                if spacing is None:
                    spacing = OxmlElement('w:spacing')
                    pPr.append(spacing)
                if is_footer:
                    spacing.set(qn('w:before'), '0')
                    spacing.set(qn('w:after'), '0')
                    spacing.set(qn('w:line'), '240')
                    spacing.set(qn('w:lineRule'), 'auto')
                else:
                    spacing.set(qn('w:after'), '240') # Ensure space below text

            if rId_map:
                _remap_rids(new_p, rId_map)

            # Force copied header/footer drawings to behave like watermark layers.
            for anchor in new_p.findall('.//' + qn('wp:anchor')):
                anchor.set('behindDoc', '1')
                anchor.set('allowOverlap', '1')

            # Force font on all runs in the copied paragraph
            for r in new_p.findall(qn('w:r')):
                rPr = r.get_or_add_rPr()
                for old_fonts in rPr.findall(qn('w:rFonts')):
                    rPr.remove(old_fonts)
                f = OxmlElement('w:rFonts')
                f.set(qn('w:ascii'), self.font_family)
                f.set(qn('w:hAnsi'), self.font_family)
                f.set(qn('w:cs'), self.font_family)
                rPr.append(f)
            target_elem.append(new_p)
    
    def _iter_all_paragraphs(self, doc: Document):
        """
        Yield every paragraph in the document,
        including paragraphs inside tables.
        """
        # Normal paragraphs
        for para in doc.paragraphs:
            yield para

        # Table paragraphs
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        yield para

    def _brand_body_paragraphs(self, doc: Document):
        heading_names = {
            'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4',
            'Heading1',  'Heading2',  'Heading3',  'Heading4',
        }

        for para in self._iter_all_paragraphs(doc):
            style_name = para.style.name if para.style else 'Normal'
            is_heading = style_name in heading_names or style_name.startswith('Heading')

            for r in para._p.findall(qn('w:r')):
                rPr = r.find(qn('w:rPr'))
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    r.insert(0, rPr)

                # ── Color: branding for headings, body_color for normal ──
                for e in rPr.findall(qn('w:color')):
                    rPr.remove(e)
                ce = OxmlElement('w:color')
                
                if is_heading:
                    # Specific Color Map
                    if '1' in style_name:
                        h_color = self.COLOR_H1
                    elif '2' in style_name:
                        h_color = self.COLOR_H2
                    elif '3' in style_name:
                        h_color = self.COLOR_H3
                    else: # H4 and below
                        h_color = self.COLOR_H4
                    ce.set(qn('w:val'), h_color)
                else:
                    ce.set(qn('w:val'), self.BODY_COLOR)
                rPr.append(ce)

                # ── Font family: always apply ──
                for e in rPr.findall(qn('w:rFonts')):
                    rPr.remove(e)
                fe = OxmlElement('w:rFonts')
                fe.set(qn('w:ascii'), self.BODY_FONT)
                fe.set(qn('w:hAnsi'), self.BODY_FONT)
                rPr.append(fe)

                # ── Remove italic everywhere ──
                for e in rPr.findall(qn('w:i')):
                    rPr.remove(e)

                if is_heading:
                    # HEADINGS: remove run-level size overrides
                    # so the STYLE-level sizes (H1=15.5pt etc.) take effect
                    for e in rPr.findall(qn('w:sz')):
                        rPr.remove(e)
                    for e in rPr.findall(qn('w:szCs')):
                        rPr.remove(e)
                else:
                    # BODY TEXT: force to BODY_SIZE
                    for e in rPr.findall(qn('w:sz')):
                        rPr.remove(e)
                    for e in rPr.findall(qn('w:szCs')):
                        rPr.remove(e)
                    s = OxmlElement('w:sz')
                    s.set(qn('w:val'), str(self.BODY_SIZE))
                    rPr.append(s)
                    sc = OxmlElement('w:szCs')
                    sc.set(qn('w:val'), str(self.BODY_SIZE))
                    rPr.append(sc)

    def _apply_page_layout_rules(self, doc: Document):
        first_h1 = True
        prev_style = None

        for para in doc.paragraphs:
            if not para.style:
                continue

            style_name = para.style.name

            if style_name in ('Heading 1', 'Heading1'):
                if not first_h1:
                    if prev_style and not prev_style.startswith('Heading'):
                        para.paragraph_format.page_break_before = True
                        pPr = para._p.find(qn('w:pPr'))
                        if pPr is None:
                            pPr = OxmlElement('w:pPr')
                            para._p.insert(0, pPr)
                        if pPr.find(qn('w:pageBreakBefore')) is None:
                            pPr.append(OxmlElement('w:pageBreakBefore'))
                first_h1 = False

            if style_name.startswith('Heading'):
                para.paragraph_format.keep_with_next = True
                pPr = para._p.find(qn('w:pPr'))
                if pPr is None:
                    pPr = OxmlElement('w:pPr')
                    para._p.insert(0, pPr)
                if pPr.find(qn('w:keepNext')) is None:
                    pPr.append(OxmlElement('w:keepNext'))

            if para.text.strip():
                prev_style = style_name
    
    def _remove_leading_empty_paragraphs(self, doc):
        while doc.paragraphs and doc.paragraphs[0].text.strip() == "":
            p = doc.paragraphs[0]._element
            p.getparent().remove(p)

    def _remove_document_protection(self, doc: Document):
        """Remove all document protection to make it editable"""
        logger.info("Removing document protection...")
        
        try:
            # Remove document protection from settings
            settings = doc.settings.element
            
            # Remove documentProtection element
            for protect in settings.findall(qn('w:documentProtection')):
                settings.remove(protect)
            
            # Remove writeProtection element (read-only)
            for write_protect in settings.findall(qn('w:writeProtection')):
                settings.remove(write_protect)
            
            # Unlock all content controls
            for content_control in doc.element.xpath('.//w:sdt'):
                sdtPr = content_control.find(qn('w:sdtPr'))
                if sdtPr is not None:
                    for lock in sdtPr.findall(qn('w:lock')):
                        sdtPr.remove(lock)
            
            logger.info("    ✓ Document is now editable")
        except Exception as e:
            logger.error(f"    ⚠ Error: {e}")
            
    def _adjust_page_margins(self, doc: Document):
        """Standardize margins and prevent header/footer overlap."""
        logger.info('Adjusting page margins...')
        for section in doc.sections:
            section.top_margin = Inches(1.25)
            section.bottom_margin = Inches(1.0)
            section.header_distance = Inches(0.5)
            section.footer_distance = Inches(0.5)

    def _fix_duplicate_media(self, doc: Document):

        """Scan document parts and relationships to find duplicate media entries."""
        logger.info("Deduplicating media...")
        try:
            pkg = doc.part.package
            media_parts = {}  # md5_hash -> part

            # 1. Collect unique media parts
            for part in list(pkg.parts):
                if 'media/' in part.partname:
                    h = get_image_hash(part.blob)

                    if h not in media_parts:
                        media_parts[h] = part

            # 2. Remap duplicate relationships
            for part in list(pkg.parts):
                if not hasattr(part, 'rels'):
                    continue

                for rel in part.rels.values():
                    if 'image' in rel.reltype:
                        target_h = get_image_hash(rel.target_part.blob)

                        if target_h in media_parts:
                            canonical_part = media_parts[target_h]

                            if rel.target_part != canonical_part:
                                # Remap to canonical image
                                rel._target = canonical_part

        except Exception as e:
            logger.warning(f"  [StyleManager] Media deduplication warning: {e}")

def _remap_rids(element, rId_map: dict):
    for attr_key in list(element.attrib.keys()):
        local = attr_key.split('}')[-1] if '}' in attr_key else attr_key
        ns    = attr_key.split('}')[0].lstrip('{') if '}' in attr_key else ''

        if ns == R_NS and local in ('embed', 'id', 'link', 'href', 'pict'):
            old_val = element.attrib[attr_key]
            if old_val in rId_map:
                element.attrib[attr_key] = rId_map[old_val]

    for child in element:
        _remap_rids(child, rId_map)
