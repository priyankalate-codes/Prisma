"""
CoverPageManager — Injects a cover page at the start of the document.

"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from copy import deepcopy
import json
import logging
import os
import requests

logger = logging.getLogger(__name__)


class CoverPageManager:

    def __init__(self, model_name: str = 'llama-3.1-8b-instant', font_family: str = 'Calibri'):
        self.model_name = model_name
        self.font_family = font_family
        self.bullet_prefixes = ('•', '-', '*', '▪', '►', '➤', '→', '○', '●', '‣', '–', '—')

    def create_cover_page(self, doc: Document) -> bool:
        """New rich cover page according to the requested design."""
        title, subtitle = self._extract_title_subtitle(doc)
        if not title:
            return False

        doc_text = self._extract_document_text(doc)
        ai_data  = self._generate_summary_with_groq(doc_text, title or "Untitled Document")
        
        # Use AI-generated title/subtitle if available, fallback to extracted ones
        display_title = ai_data.get('title', title)
        display_subtitle = ai_data.get('subtitle', subtitle)
        abstract = ai_data.get('abstract', '')
        coverage = ai_data.get('coverage', [])
        prev_v   = ai_data.get('prev_version', 'v0.9 Draft')
        phase    = ai_data.get('phase', 'Phase 1 - Initial Specs')

        # We will build all elements and then insert them sequentially at the TOP
        elements_to_insert = []

        # 1. Product Specification Header (Ultra-Tight Spacing)
        p_header = self._create_styled_p(doc, "PRODUCT SPECIFICATION", size=10, bold=True, color='2F5496')
        p_header.paragraph_format.space_before = Pt(0)
        p_header.paragraph_format.space_after = Pt(2)
        elements_to_insert.append(p_header._element)

        # 2. Main Title (Ultra-Tight Spacing)
        p_title = self._create_styled_p(doc, display_title, size=44, bold=True, color='2E2E2E')
        p_title.paragraph_format.space_before = Pt(0)
        p_title.paragraph_format.space_after = Pt(0)
        elements_to_insert.append(p_title._element)

        # 3. Subtitle (Ultra-Tight Spacing)
        if display_subtitle:
            p_sub = self._create_styled_p(doc, display_subtitle, size=24, color='7F7F7F')
            p_sub.paragraph_format.space_before = Pt(0)
            p_sub.paragraph_format.space_after = Pt(4)
            elements_to_insert.append(p_sub._element)

        # 4. First Blue Horizontal Line
        elements_to_insert.append(self._create_horizontal_line(doc, color='2F5496', thickness=12))

        # 5. Abstract text (Tightest padding)
        if abstract:
            p_abs = self._create_styled_p(doc, abstract, size=11, color='595959')
            p_abs.paragraph_format.space_before = Pt(12)
            p_abs.paragraph_format.space_after = Pt(12)
            elements_to_insert.append(p_abs._element)

        # 6. Coverage Labels & Pills (Tightest padding)
        if coverage:
            label = self._create_styled_p(doc, "COVERAGE", size=9, bold=True, color='7F7F7F')
            label.paragraph_format.space_after = Pt(2)
            elements_to_insert.append(label._element)
            
            p_tags = self._create_coverage_tags(doc, coverage)
            p_tags.paragraph_format.space_after = Pt(12)
            elements_to_insert.append(p_tags._element)

        # 7. Separator above Table
        elements_to_insert.append(self._create_horizontal_line(doc, color='E9EEF4', thickness=4))

        # 8. Metadata Table (8 rows)
        elements_to_insert.append(self._create_metadata_table(doc, display_title, display_subtitle, prev_v, phase))

        # 9. Separator below Table
        elements_to_insert.append(self._create_horizontal_line(doc, color='2F5496', thickness=8))

        # 10. Footer Italic Summary
        classification = "Confidential — Internal & Authorised Partners Only"
        elements_to_insert.append(self._create_footer_summary(doc, display_title, classification))
        
        # 11. Final Decorative Bottom Line
        elements_to_insert.append(self._create_horizontal_line(doc, color='2F5496', thickness=4))

        # 12. Section Break
        elements_to_insert.append(self._make_cover_section_break(doc))

        # Insert them into the body
        body = doc.element.body
        for i, elem in enumerate(elements_to_insert):
            body.insert(i, elem)

        return True

    def _create_styled_p(self, doc: Document, text: str, size: int,
                        bold: bool = False, italic: bool = False, color: str = '000000'):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        
        # Color parsing (hex to RGB)
        if color.startswith('#'): color = color[1:]
        r.font.color.rgb = RGBColor(
            int(color[0:2], 16),
            int(color[2:4], 16),
            int(color[4:6], 16)
        )
        r.font.name = self.font_family
        
        # Remove from auto-append
        p._element.getparent().remove(p._element)
        return p

    def _create_horizontal_line(self, doc: Document, color: str, thickness: int = 6):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(thickness))
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), color)
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        p._element.getparent().remove(p._element)
        return p._element

    def _create_coverage_tags(self, doc: Document, topics: list):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(20)
        
        for topic in topics:
            r = p.add_run(topic.upper())
            r.font.size = Pt(9)
            r.font.bold = True
            r.font.name = self.font_family
            r.font.color.rgb = RGBColor(255, 255, 255)
            
            # Apply Shading (The "Pill" effect)
            rPr = r._element.get_or_add_rPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '2563EB') # Dynamic Blue
            rPr.append(shd)
            
            # Spacer after tag
            p.add_run("   ")
        
        p._element.getparent().remove(p._element)
        return p

    def _create_metadata_table(self, doc: Document, title: str, subtitle: str, prev_v: str = "TBD", phase: str = "Planning"):
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Table Grid'
        
        # 1. Update Table Borders
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        tblBorders = OxmlElement('w:tblBorders')
        
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '12')
        left.set(qn('w:color'), '5D8BF0')
        tblBorders.append(left)
        
        for b in ('top', 'bottom', 'right', 'insideH', 'insideV'):
            edge = OxmlElement(f'w:{b}')
            edge.set(qn('w:val'), 'single')
            edge.set(qn('w:sz'), '4')
            edge.set(qn('w:color'), 'E9EEF4')
            tblBorders.append(edge)
            
        tblPr.append(tblBorders)

        date_str = datetime.now().strftime('%B %Y')
        metadata = [
            ("Document", f"{title} — {subtitle if subtitle else 'Project Specification'}"),
            ("Version", "v1.0"),
            ("Status", "In Review — Active Prototype"),
            ("Date", date_str),
            ("Audience", "Production Leads, Tech Stakeholders"),
            ("Classification", "Confidential — Internal & Authorised Partners Only"),
            ("Previous Version", prev_v),
            ("Phase Coverage", phase)
        ]
        
        row_colors = ['F1F5F9', 'F8FAFC']
        
        for i, (label, value) in enumerate(metadata):
            current_row_color = row_colors[i % 2]
            
            cell_label = table.cell(i, 0)
            cell_label.text = label
            run_lbl = cell_label.paragraphs[0].runs[0]
            run_lbl.font.bold = True
            run_lbl.font.size = Pt(10)
            run_lbl.font.name = self.font_family
            run_lbl.font.color.rgb = RGBColor(89, 89, 89)
            
            cell_val = table.cell(i, 1)
            cell_val.text = value
            run_val = cell_val.paragraphs[0].runs[0]
            run_val.font.size = Pt(10)
            run_val.font.name = self.font_family
            run_val.font.color.rgb = RGBColor(46, 46, 46)
            
            for cell in (cell_label, cell_val):
                tcPr = cell._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), current_row_color)
                tcPr.append(shd)

        table._element.getparent().remove(table._element)
        return table._element

    def _create_footer_summary(self, doc: Document, title: str, classification: str):
        date_str = datetime.now().strftime('%B %Y')
        summary_text = f"{classification}  ·  {title} Internal Documentation  ·  {date_str}"
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(summary_text)
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.name = self.font_family
        r.font.color.rgb = RGBColor(127, 127, 127)
        
        p._element.getparent().remove(p._element)
        return p._element

    def _make_cover_section_break(self, doc: Document):
        """
        Paragraph containing a next-page section break with vertical centering.
        Copies header/footer references from the template so the cover page
        gets the company letterhead too.
        """
        p      = OxmlElement('w:p')
        pPr    = OxmlElement('w:pPr')
        sectPr = OxmlElement('w:sectPr')

        # Copy header/footer references and page size/margins from template
        template_sectPr = doc.sections[0]._sectPr
        for tag in ('w:headerReference', 'w:footerReference', 'w:pgSz', 'w:pgMar'):
            for elem in template_sectPr.findall(qn(tag)):
                sectPr.append(deepcopy(elem))

        # Next page break so content starts on page 2
        sectType = OxmlElement('w:type')
        sectType.set(qn('w:val'), 'nextPage')
        sectPr.append(sectType)

        pPr.append(sectPr)
        p.append(pPr)
        return p

    def _force_para_center_xml(self, para):
        """Write center alignment directly to XML so it cannot be overridden by styles."""
        pPr = para._element.get_or_add_pPr()
        for existing in pPr.findall(qn('w:jc')):
            pPr.remove(existing)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # AI SUMMARY + TEXT EXTRACTION
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

    def _generate_summary_with_groq(self, doc_text: str, title_hint: str) -> dict:
        """Generate concise titles, summaries, and coverage topics as JSON using Groq."""
        if len(doc_text) > 8000:
            doc_text = doc_text[:8000]

        api_key = os.getenv("GROQ_API")
        if not api_key:
            logger.warning("GROQ_API key not found, using fallback cover-page metadata")
            return self._default_cover_metadata()

        prompt = f"""
        Analyze the following document. The original heading is "{title_hint}".
        
        Generate the following metadata:
        1. 'title': A strictly 2-word summary of the main subject.
        2. 'subtitle': A strictly 3-4 word description of the document's scope.
        3. 'abstract': A 1-2 sentence executive summary of the document's purpose.
        4. 'coverage': A list of 5-8 key modules or features mentioned.
        5. 'prev_version': A version number (e.g. v0.9) and very brief delta (if found, else default v0.1).
        6. 'phase': A 3-5 word description of the current phase (e.g. Phase 1 MVP + Feedback).

        Return strictly as a JSON object with keys "title", "subtitle", "abstract", "coverage", "prev_version", "phase".
        Example:
        {{
          "title": "AI Trends",
          "subtitle": "Future Intelligence Analysis 2026",
          "abstract": "Detailed technical specification for the TableFlow platform.",
          "coverage": ["Auth", "Dashboard", "API"],
          "prev_version": "v1.1 - Added Core Logic",
          "phase": "Phase 1 - Initial Deployment"
        }}

        Document Content:
        {doc_text}
        """
        
        try:
            response = requests.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": self.model_name,
                    "temperature": 0.1,
                    "messages": [
                        {
                            "role": "system",
                            "content": "Return only valid JSON with the exact requested keys."
                        },
                        {
                            "role": "user",
                            "content": prompt
                        }
                    ]
                },
                timeout=30
            )
            response.raise_for_status()
            content = response.json().get('choices', [{}])[0].get('message', {}).get('content', '').strip()
            if '```json' in content:
                content = content.split('```json')[1].split('```')[0].strip()
            elif '{' in content:
                content = content[content.find('{'):content.rfind('}')+1]
            
            return json.loads(content)
        except Exception as e:
            logger.error(f"Groq cover-page extraction failed: {e}")
            return self._default_cover_metadata()

    def _default_cover_metadata(self) -> dict:
        return {
            "title": "Document Overview",
            "subtitle": "Structural Analysis Report",
            "abstract": "Comprehensive documentation highlighting key project requirements.",
            "coverage": ["System Architecture", "Functional Requirements", "Interface Design"],
            "prev_version": "v0.9 - Structural Draft",
            "phase": "Draft Exploration Phase"
        }

    def _extract_title_subtitle(self, doc: Document) -> tuple:
        """
        Scan the first 15 paragraphs for the title and subtitle.
        Skips empty lines, TOC headings, and bullet points.
        Returns (title, subtitle) — either can be None.
        """
        candidates = []
        for para in doc.paragraphs[:15]:
            text = para.text.strip()
            if not text:
                continue
            if text.upper() == 'TABLE OF CONTENTS':
                continue
            if any(text.startswith(prefix) for prefix in self.bullet_prefixes):
                continue
            if 2 <= len(text.split()) <= 25:
                candidates.append(text)
            if len(candidates) >= 2:
                break
        return (
            candidates[0] if candidates else None,
            candidates[1] if len(candidates) > 1 else None
        )

    def _extract_document_text(self, doc: Document) -> str:
        """Return all non-empty paragraph text joined by newlines."""
        return '\n'.join(p.text.strip() for p in doc.paragraphs if p.text.strip())
