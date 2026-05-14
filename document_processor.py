"""
DocumentProcessor — Signal-Based Classifier

"""

import os
import re
import traceback
import logging
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
from utils.style_manager import StyleManager

# Configure logging
logger = logging.getLogger(__name__)
try:
    from utils.adobe_helper import adobe_pdf_extract
except ImportError:
    adobe_pdf_extract = None
from utils.toc_manager import TocManager
from utils.cover_page_manager import CoverPageManager
from config import ADOBE_CLIENT_ID, ADOBE_CLIENT_SECRET

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


class DocumentProcessor:
    def __init__(self, template_path: str, font_family: str = 'Calibri', font_size: int = 11, 
                 include_cover: bool = False, include_toc: bool = False):
        self.template_path = template_path
        self.style_manager = StyleManager(template_path, font_family, font_size)
        self.include_cover = include_cover
        self.include_toc = include_toc
        self.toc_manager = TocManager(font_family=font_family)
        self.cover_manager = CoverPageManager(font_family=font_family)

    def _parse_style_list_level(self, style_name: str) -> int:
        if not style_name:
            return 0
        match = re.search(r'List\s+(?:Bullet|Number)\s*(\d+)?', style_name, re.IGNORECASE)
        if not match:
            return 0
        return max(0, int(match.group(1) or '1') - 1)

    def _infer_indent_level(self, left_indent) -> int:
        if left_indent is None:
            return 0
        try:
            inches = left_indent.inches
            if inches <= 0.30:
                return 0
            return max(0, int(round((inches - 0.25) / 0.22)))
        except Exception:
            return 0

    def _get_list_style_name(self, numbered: bool, level: int) -> str:
        family = 'List Number' if numbered else 'List Bullet'
        style_level = min(max(0, int(level or 0)), 2)
        if style_level == 0:
            return family
        return f'{family} {style_level + 1}'

    def _strip_leading_list_marker(self, text: str) -> str:
        """Remove source list markers before applying a Word list style."""
        if not text:
            return text
        return re.sub(
            r'^(?:[-–—]|\d+[\.\)]|[☐☑•▪►■□●○◦⁃‣])\s*',
            '',
            text.strip()
        ).strip()

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # PUBLIC API
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

    def universal_extract(self, input_path: str, output_path: str) -> dict:
        """Entry point for uploaded files. Routes by extension."""
        try:
            ext = os.path.splitext(input_path)[1].lower()
            if ext == '.pdf':
                result = self._pipeline_pdf(input_path, output_path)
            elif ext == '.docx':
                result = self._pipeline_docx(input_path, output_path)
            elif ext == '.txt':
                result = self._pipeline_txt(input_path, output_path)
            else:
                return {'success': False, 'error': f'Unsupported file type: {ext}'}

            if result.get('success') and self.include_toc:
                self.refresh_saved_docx(output_path)
            return result
        except Exception as e:
            traceback.print_exc()
            return {'success': False, 'error': str(e)}

    def refresh_saved_docx(self, output_path: str):
        """Refresh TOC/page fields in a saved DOCX when Word automation is available."""
        if self.include_toc:
            self.toc_manager.refresh_toc_page_numbers(output_path)

    def html_to_docx(self, html: str) -> Document:
        doc = Document(self.template_path)
        soup = BeautifulSoup(html, 'html.parser')

        # If there are NO block-level tags like p, h1, div, etc., 
        # then this is likely plain text from the 'Edit' box.
        # In this case, we use the Smart Classifier to detect headings.
        has_block_tags = bool(soup.find(['p', 'h1', 'h2', 'h3', 'h4', 'div', 'ul', 'ol', 'table']))

        if not has_block_tags:
            logger.info("Detected plain text input (likely Edit), using Smart Classifier.")
            lines = html.splitlines()
            raw_lines = []
            for line in lines:
                t = line.strip()
                if t:
                    # Mock the signals expected by the classifier
                    raw_lines.append({
                        'text': t,
                        'style': 'Normal',
                        'is_bold': False,      # Plain text doesn't have bold signals
                        'run_size': 0,
                        'num_id': 0,
                        'num_lvl': 0
                    })
            self._h4_counters = {}
            self._build_from_signals(raw_lines, doc)
        else:
            logger.info(f"Detected HTML input (length: {len(html)}), using tag-based parser.")
            processed_elements = set()

            def get_block_container(container):
                return container._parent if isinstance(container, Paragraph) else container

            def add_block_paragraph(container, style=None):
                block_container = get_block_container(container)
                if style:
                    return block_container.add_paragraph(style=style)
                return block_container.add_paragraph()

            def add_block_heading(container, level: int):
                block_container = get_block_container(container)
                return block_container.add_heading('', level=level)

            def process_node(node, container):
                if getattr(node, 'name', None) and id(node) in processed_elements:
                    return
                if not getattr(node, 'name', None):
                    clean_text = str(node).strip()
                    if clean_text:
                        if isinstance(container, Paragraph):
                            style_name = container.style.name if container.style else ''
                            if 'Bullet' in style_name or 'Number' in style_name:
                                clean_text = self._strip_leading_list_marker(clean_text)
                            container.add_run(clean_text)
                        else: container.add_paragraph(clean_text)
                    return

                if node.name in ('h1', 'h2', 'h3', 'h4'):
                    level = int(node.name[1])
                    para = add_block_heading(container, level)
                    for child in node.contents: process_node(child, para)
                    processed_elements.add(id(node))
                elif node.name == 'p':
                    para = add_block_paragraph(container)
                    for child in node.contents: process_node(child, para)
                    processed_elements.add(id(node))
                elif node.name == 'table':
                    self._add_html_table(get_block_container(container), node)
                    processed_elements.add(id(node))
                elif node.name in ('ul', 'ol'):
                    for li in node.find_all('li', recursive=False):
                        if node.name == 'ol':
                            para = add_block_paragraph(container, style='List Number')
                        else:
                            para = add_block_paragraph(container, style='List Bullet')
                        for child in li.contents:
                            process_node(child, para)
                        self.style_manager.apply_list_styling(para)
                    processed_elements.add(id(node))
                elif node.name in ('b', 'strong', 'i', 'em', 'u', 'span'):
                    if isinstance(container, Paragraph):
                        run = container.add_run()
                        if node.name in ('b', 'strong'): run.bold = True
                        if node.name in ('i', 'em'): run.italic = True
                        if node.name == 'u': run.underline = True
                        for child in node.contents:
                            if not getattr(child, 'name', None): run.text += str(child)
                            else: process_node(child, container)
                    else:
                        para = add_block_paragraph(container)
                        process_node(node, para)
                    processed_elements.add(id(node))
                elif node.name == 'div':
                    block_children = node.find(['p', 'h1', 'h2', 'h3', 'h4', 'ul', 'ol', 'table', 'div'])
                    if not block_children:
                        para = add_block_paragraph(container)
                        for child in node.contents: process_node(child, para)
                        processed_elements.add(id(node))
                    else:
                        for child in node.contents: process_node(child, container)
                else:
                    for child in node.contents: process_node(child, container)

            for child in soup.contents:
                process_node(child, doc)

        styled_doc = self.style_manager.apply_template_styles(doc)
        self._apply_final_features(styled_doc)
        return styled_doc

    def _apply_final_features(self, doc: Document):
        """Adds Cover Page and Table of Contents if requested."""
        if self.include_cover:
            self.cover_manager.create_cover_page(doc)

        if self.include_toc:
            self.toc_manager.insert_toc(doc)

    def _add_html_table(self, doc: Document, element):
        """
        Robust HTML table parser.
        Handles nested tables, merged cells, and complex HTML structures.
        """
        try:
            # Find the table element
            table = element
            
            # Create Word table with estimated dimensions
            rows = table.find_all('tr')
            num_rows = len(rows)
            
            # Estimate columns by finding max cells in any row
            num_cols = 0
            for row in rows:
                cells = row.find_all(['td', 'th'])
                num_cols = max(num_cols, len(cells))
            
            if num_rows == 0 or num_cols == 0:
                return
            
            # Create Word table
            word_table = doc.add_table(rows=num_rows, cols=num_cols)
            word_table.style = 'Table Grid'
            
            # Process each row
            for i, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                
                # Track merged cells
                col_offset = 0
                
                for cell in cells:
                    # Get cell text
                    cell_text = ''.join(cell.get_text(separator=' ', strip=True))
                    
                    # Get cell attributes
                    rowspan = int(cell.get('rowspan', 1))
                    colspan = int(cell.get('colspan', 1))
                    
                    # Find the target cell in Word table
                    target_row = i
                    target_col = col_offset
                    
                    # Adjust for previous merged cells
                    while target_row < num_rows and target_col < num_cols:
                        if word_table.cell(target_row, target_col).text.strip() == '':
                            break
                        target_col += 1
                    
                    if target_row >= num_rows or target_col >= num_cols:
                        continue
                    
                    # Merge cells if needed
                    if rowspan > 1 or colspan > 1:
                        word_table.cell(target_row, target_col).merge(
                            word_table.cell(target_row + rowspan - 1, target_col + colspan - 1)
                        )
                    
                    # Add content to cell
                    word_cell = word_table.cell(target_row, target_col)
                    word_cell.text = cell_text
                    
                    # Apply basic styling
                    if cell.name == 'th':
                        word_cell.paragraphs[0].runs[0].bold = True
                        word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Move to next column
                    col_offset += colspan
                    
        except Exception as e:
            logger.error(f"Error parsing HTML table: {e}")
            # Fallback: add as plain text
            doc.add_paragraph(str(element))
    
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # DOCX PIPELINE
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

    def _pipeline_docx(self, path: str, out: str) -> dict:
        source_doc = Document(path)
        new_doc    = Document()
        self._h4_counters = {}  

        # ── STEP 1: EXTRACT (Preserving Order) ──────────────────────────
        signals = self._extract_docx_signals(source_doc)
        signals = self._stitch_number_fragments(signals)
        
        # ── STEP 2: JOIN ─────────────────────────────────────────────────
        # Word COM PDF→DOCX conversion fragments numbered items across paragraphs:
        
        joined = []
        i = 0
        while i < len(signals):
            item = signals[i]
            t    = item['text'].strip()

            # Lone digit — attempt forward stitch
            if re.match(r'^\d+$', t) and i + 1 < len(signals):
                next_t = signals[i + 1]['text'].strip()

                if re.match(r'^[\.\)]\s+\S', next_t):
                    # "1" + ". Customer opens..." → "1. Customer opens..."
                    # Normalize both "." and ")" separators to ". "
                    normalized = re.sub(r'^[\.\)]\s*', '. ', next_t)
                    item = dict(item)
                    item['text'] = t + normalized
                    joined.append(item)
                    i += 2
                    continue

                elif next_t in ('.', ')') and i + 2 < len(signals):
                    # "3" + ")" + ". System validates..." → "3. System validates..."
                    rest       = signals[i + 2]['text'].strip()
                    rest_clean = re.sub(r'^[\.\)]\s*', '', rest)
                    item = dict(item)
                    item['text'] = t + '. ' + rest_clean
                    joined.append(item)
                    i += 3
                    continue

                elif re.match(r'^\d+$', next_t):
                    # Two consecutive lone digits — skip this one
                    i += 1
                    continue

                else:
                    # Lone digit with no joinable continuation — keep as placeholder
                    item = dict(item)
                    item['text'] = t + '.'
                    joined.append(item)
                    i += 1
                    continue

            # Lone punctuation — append to previous only if it isn't already a complete numbered item
            elif joined and len(t) == 1 and t in '.)':
                if not re.match(r'^\d+[\.\)]', joined[-1]['text']):
                    joined[-1]['text'] += t
                i += 1
                continue

            # Fragment like ". Customer scans QR on table"
            elif joined and re.match(r'^[\.\)]\s+\S', t):
                prev = joined[-1]['text']
                if not re.match(r'^\d+[\.\)]\s+\S', prev):
                    # Previous is not yet a complete numbered item — stitch
                    joined[-1]['text'] += t
                    i += 1
                    continue
                else:
                    # Previous is already a complete numbered item — this is a
                    # continuation bullet from the same step; strip leading dot
                    item = dict(item)
                    item['text'] = re.sub(r'^[\.\)]\s*', '', t).strip()
                    if item['text']:
                        joined.append(item)
                    i += 1
                    continue

            # Lone "-" paragraph — prefix the NEXT line as a dash bullet
            elif t == '-':
                if i + 1 < len(signals):
                    signals[i + 1]['text'] = '- ' + signals[i + 1]['text'].lstrip('- ')
                i += 1
                continue

            if t:
                joined.append(item)
            i += 1

        signals = joined

        # ── STEPS 3 & 4: CLASSIFY + BUILD ──────────────────────────────────
        self._build_from_signals(signals, new_doc)

        branded_doc = self.style_manager.apply_template_styles(new_doc)
        self._apply_final_features(branded_doc)
        
        branded_doc.save(out)
        return {'success': True, 'paragraphs': len(new_doc.paragraphs)}

    def _pipeline_txt(self, path: str, out: str) -> dict:
        with open(path, 'r', encoding='utf-8', errors='replace') as handle:
            text = handle.read()

        doc = self.html_to_docx(text)
        doc.save(out)
        return {'success': True, 'paragraphs': len(doc.paragraphs)}

    def _stitch_number_fragments(self, signals):
        """Join split list markers like '1.' + 'Item text' before classification."""
        stitched = []
        i = 0

        while i < len(signals):
            item = signals[i]
            text = item.get('text', '').strip()
            next_item = signals[i + 1] if i + 1 < len(signals) else None
            next_text = next_item.get('text', '').strip() if next_item else ''

            current_number = re.match(r'^(\d+)(?:[\.\)])?$', text)
            next_is_number_only = bool(re.match(r'^\d+(?:[\.\)])?$', next_text))

            if current_number and next_text and not next_is_number_only:
                merged = dict(item)
                merged['text'] = f"{current_number.group(1)}. {next_text.lstrip('. )')}".strip()
                stitched.append(merged)
                i += 2
                continue

            stitched.append(item)
            i += 1

        return stitched

    def _build_from_signals(self, signals, doc):
        """Standard processing logic shared by initial files and text edits."""
        signals = self._stitch_number_fragments(signals)
        expecting_list = False
        is_first_non_empty = True
        last_item_type = None
        list_positions = {}

        def get_insertion_point():
            """Returns the standard place to insert Word elements (before sectPr)."""
            body = doc.element.body
            sectPr = body.find(qn('w:sectPr'))
            return sectPr

        def touch_list_level(level: int):
            list_positions[level] = list_positions.get(level, 0) + 1
            for key in list(list_positions.keys()):
                if key > level:
                    del list_positions[key]

        def clear_list_context():
            list_positions.clear()

        for idx, item in enumerate(signals):
            item_type = item.get('type')
            
            # --- Table Handling ---
            if item_type == 'table':
                # Add spacer if following another table
                if last_item_type == 'table':
                    doc.add_paragraph("") 
                
                # Create a "hollow" table to copy the OXML into
                # Or just insert the OXML directly into the body at the right spot
                table_node = item['node']
                target = get_insertion_point()
                if target is not None:
                    target.addprevious(table_node)
                else:
                    doc.element.body.append(table_node)
                
                last_item_type = 'table'
                continue

            text     = item['text'].replace('Ł', '').strip()
            style    = item.get('style', 'Normal')
            is_bold  = item.get('is_bold', False)
            run_size = item.get('run_size', 0)
            num_id   = item.get('num_id', 0)
            num_lvl  = item.get('num_lvl', 0)
            list_level = item.get('list_level', num_lvl)
            is_numbered_list_item = num_id > 0

            num_fmt  = item.get('num_fmt', '')
            is_list_bullet_style = 'Bullet' in style
            is_list_number_style = 'Number' in style
            is_bullet = num_fmt == 'bullet' or is_list_bullet_style

            if not text: continue

            next_text = ''
            next_item_type = None
            for future in signals[idx + 1:]:
                future_type = future.get('type')
                if future_type == 'table':
                    next_item_type = future_type
                    break
                candidate = future.get('text', '').replace('Å', '').strip()
                if candidate:
                    next_text = candidate
                    next_item_type = future_type
                    break
            
            # --- Signals ---
            is_section_label = text.endswith(':') and len(text.split()) <= 8
            word_count = len(text.split())
            explicit_number_match = re.match(r'^(\d+)[\.\)]\s+(.+)$', text)
            numbered_heading_match = re.match(r'^((?:\d+\.)+\d+|\d+[\.\)])\s+(.+)$', text)
            is_explicit_numbered_item = bool(explicit_number_match)
            is_numbered_heading = bool(numbered_heading_match)
            looks_like_heading_text = word_count <= 12 and not text.endswith('.')
            is_structural_heading = bool(
                re.match(r'^(section|major section|chapter|part|appendix|workstream)\s+\S+', text, re.IGNORECASE)
            ) and word_count <= 6
            is_short_numbered_heading = bool(
                numbered_heading_match
                and len(numbered_heading_match.group(2).split()) <= 6
                and not numbered_heading_match.group(2).strip().endswith('.')
            )
            is_large_heading_signal = (
                run_size >= 34 and looks_like_heading_text and word_count <= 8
            )
            is_medium_heading_signal = (
                run_size >= 28 and looks_like_heading_text and word_count <= 8
            )
            is_contextual_heading_candidate = (
                style == 'Normal'
                and run_size == 0
                and not is_bold
                and not text.endswith((':', '.'))
                and word_count <= 5
                and bool(next_text)
                and next_item_type != 'table'
                and len(next_text.split()) >= 6
                and next_text.endswith('.')
                and not bool(numbered_heading_match)
            )
            # Broadened markers: check for bullet characters and custom markers
            is_explicit_bullet  = text.startswith(('☐', '☑', '•', '▪', '►', '■', '□', '●', '○', '◦', '⁃', '‣'))
            is_dash_bullet      = bool(re.match(r'^[-–—]\s+\S', text))
            is_list_like = (
                is_numbered_list_item
                or is_list_bullet_style
                or is_list_number_style
                or is_explicit_bullet
                or is_dash_bullet
                or list_level > 0
            )

            # --- Classifier logic (H1-H4) ---
            if 'Heading 1' in style: 
                doc.add_heading(text, level=1)
                clear_list_context()
            elif 'Heading 2' in style:
                doc.add_heading(text, level=2)
                clear_list_context()
            elif 'Heading 3' in style:
                doc.add_heading(text, level=3)
                clear_list_context()
            elif is_first_non_empty and len(text.split()) <= 15 and not text.endswith('.') and not is_list_like:
                doc.add_heading(text, level=1)
                is_first_non_empty = False
                last_item_type = 'text'
                continue
            elif not is_list_like and (
                is_structural_heading
                or is_large_heading_signal
                or is_contextual_heading_candidate
            ):
                doc.add_heading(text, level=1)
                clear_list_context()
            elif not is_list_like and (
                not expecting_list
                and not is_section_label
                and (
                    is_short_numbered_heading
                    or is_medium_heading_signal
                )
            ):
                doc.add_heading(text, level=2)
                clear_list_context()
            elif not is_list_like and (
                is_section_label or (is_bold and word_count <= 8 and not text.endswith('.'))
            ):
                doc.add_heading(text, level=3)
                clear_list_context()
            elif is_explicit_numbered_item and not is_numbered_list_item and not is_list_number_style:
                clean = explicit_number_match.group(2).strip().rstrip('-').strip()
                if clean:
                    touch_list_level(list_level)
                    para = doc.add_paragraph(style=self._get_list_style_name(True, list_level))
                    para.add_run(clean)
                    self.style_manager.apply_list_styling(para, level=list_level)
            elif (
                is_bullet
                or is_list_number_style
                or is_numbered_list_item
                or is_explicit_bullet
                or is_dash_bullet
                or list_level > 0
                or expecting_list
            ):
                # Clean up marker and trailing dash artifacts
                clean = re.sub(r'^[-–—☐☑•▪►■□●○◦⁃‣]\s*', '', text).strip().rstrip('-').strip()
                if clean:
                    try:
                        # Default contextual lists to bullets unless the item is
                        # explicitly numbered or carries a clear numbered-list style.
                        touch_list_level(list_level)
                        should_use_numbering = (
                            (is_numbered_list_item or is_list_number_style)
                            and not is_bullet
                        )
                        if should_use_numbering:
                            para = doc.add_paragraph(style=self._get_list_style_name(True, list_level))
                            para.add_run(clean)
                        elif is_list_bullet_style or is_list_number_style:
                            para = doc.add_paragraph(clean, style=style)
                        else:
                            style_target = self._get_list_style_name(should_use_numbering, list_level)
                            para = doc.add_paragraph(clean, style=style_target)

                        self.style_manager.apply_list_styling(para, level=list_level)
                    except: 
                        para = doc.add_paragraph(clean)
            else:
                doc.add_paragraph(text)
                clear_list_context()

            is_first_non_empty = False
            
            # Update list expectation state
            if is_section_label:
                expecting_list = True
            elif text.endswith('.') or len(text) > 120:
                expecting_list = False
                
            last_item_type = 'text'

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # PDF PIPELINE
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

    def _pipeline_pdf(self, path: str, out: str) -> dict:
        """Adobe Extract API → semantic JSON elements → branded DOCX."""
        try:
            if adobe_pdf_extract is None:
                logger.error("Adobe extractor is unavailable. Check the Adobe SDK installation.")
                return {'success': False, 'error': 'Adobe service unavailable'}

            if not ADOBE_CLIENT_ID or not ADOBE_CLIENT_SECRET:
                logger.error("Adobe credentials are missing.")
                return {'success': False, 'error': 'Adobe credentials are missing'}

            elements = adobe_pdf_extract(path, ADOBE_CLIENT_ID, ADOBE_CLIENT_SECRET)
            if elements is None:
                logger.error("Adobe returned no data. Check your Client ID and BASE_UPLOAD_URL.")
                return {'success': False, 'error': 'Adobe service unavailable'}

            doc         = self._build_from_adobe_json(elements)
            branded_doc = self.style_manager.apply_template_styles(doc)
            
            # Apply optional features
            self._apply_final_features(branded_doc)

            branded_doc.save(out)
            return {'success': True}
        except Exception as e:
            logger.error(f'Adobe API error: {e}')
            return {'success': False, 'error': str(e)}

    def _build_from_adobe_json(self, elements: list) -> Document:
        """Convert Adobe Extract API semantic elements into a structured DOCX."""
        doc            = Document(self.template_path)
        expecting_list = False

        for el in elements:
            path = el.get('Path', '')
            text = el.get('Text', '').strip()

            if not text and 'Table' not in path:
                continue

            # Adobe separates bullet labels ("/Lbl") from content ("/LBody").
            # Skip labels — LBody carries the full text we need.
            if '/Lbl' in path:
                continue

            text = text.replace('Ł', '').strip()

            is_section_label = text.endswith(':') and len(text.split()) <= 6
            is_explicit_bullet = text.startswith(('-', '☐', '•', '▪', '➤', '■'))

            # Update list context
            if is_section_label:
                expecting_list = True
            elif text.endswith('.') or len(text) > 120 or re.search(r'/H\d', path):
                expecting_list = False

            if '/Title' in path:
                doc.add_heading(text, level=1)
            elif re.search(r'/H1', path):
                doc.add_heading(text, level=1)
            elif re.search(r'/H2', path):
                doc.add_heading(text, level=2)
            elif re.search(r'/H3', path) or is_section_label:
                doc.add_heading(text, level=3)
            elif re.search(r'/H4', path) or (('/LI' in path or '/LBody' in path) and re.match(r'^\d+\.', text)):
                # If it's a numbered item like "1. Task", promote to H4
                doc.add_heading(text, level=4)
            elif '/LI' in path or '/LBody' in path or is_explicit_bullet or expecting_list:
                clean = re.sub(r'^[-\[\]☐•▪➤■]\s*', '', text)
                doc.add_paragraph(clean, style='List Bullet')
            elif 'Table' in path:
                pass  # table reconstruction not yet implemented
            else:
                doc.add_paragraph(text)

        return doc

    def _extract_docx_signals(self, doc):
        """Iterates through the document body to maintain the correct order of elements."""
        signals = []
        
        # doc.element.body contains all elements (paragraphs and tables) in order
        for element in doc.element.body:
            # 1. PARAGRAPH (CT_P)
            if isinstance(element, CT_P):
                # Word-COM conversion stripping / Fallback check
                is_fallback = False
                parent = element.getparent()
                while parent is not None:
                    if (parent.tag.split('}')[-1] if '}' in parent.tag else parent.tag) == 'Fallback':
                        is_fallback = True; break
                    parent = parent.getparent()
                if is_fallback: continue

                para       = Paragraph(element, doc)
                text_full  = para.text.strip()
                if not text_full: continue

                style_name = para.style.name if para.style else 'Normal'
                list_level = self._parse_style_list_level(style_name)
                inferred_indent_level = self._infer_indent_level(para.paragraph_format.left_indent)
                
                # Word list-numbering (XML Signals)
                pPr      = element.find(qn('w:pPr'))
                numPr    = pPr.find(qn('w:numPr'))   if pPr   is not None else None
                numId_el = numPr.find(qn('w:numId')) if numPr is not None else None
                ilvl_el  = numPr.find(qn('w:ilvl'))  if numPr is not None else None
                num_id   = int(numId_el.get(qn('w:val'), 0)) if numId_el is not None else 0
                num_lvl  = int(ilvl_el.get(qn('w:val'),  0)) if ilvl_el  is not None else 0
                if num_id > 0:
                    list_level = num_lvl
                else:
                    list_level = max(list_level, inferred_indent_level)

                # Run-level signals: bold and size
                run_sizes, is_bold = [], False
                for r in element.findall(qn('w:r')):
                    rPr = r.find(qn('w:rPr'))
                    if rPr is not None:
                        sz = rPr.find(qn('w:sz'))
                        if sz is not None:
                            try: run_sizes.append(int(sz.get(qn('w:val'))))
                            except: pass
                        if self._xml_flag_enabled(rPr, 'w:b'):
                            is_bold = True
                run_size = max(run_sizes) if run_sizes else 0

                # Detect format (bullet vs number)
                num_fmt = 'decimal'
                if numId_el is not None:
                    try:
                        num = doc.part.numbering_part.numbering_definitions._numbering.num_having_numId(num_id)
                        abstractNum = num.abstractNum
                        for lvl in abstractNum.findall(qn('w:lvl')):
                            if int(lvl.get(qn('w:ilvl'))) == num_lvl:
                                numFmt_el = lvl.find(qn('w:numFmt'))
                                if numFmt_el is not None:
                                    num_fmt = numFmt_el.get(qn('w:val'))
                                break
                    except Exception:
                        pass

                signals.append({
                    'type': 'text',
                    'text': text_full,
                    'style': style_name,
                    # Extended metadata for classifier
                    'is_bold': is_bold,
                    'run_size': run_size,
                    'num_id': num_id,
                    'num_lvl': num_lvl,
                    'num_fmt': num_fmt,
                    'list_level': list_level
                })
            
            # 2. TABLE (CT_Tbl)
            elif isinstance(element, CT_Tbl):
                # Use current table-to-html processor
                table_html = self._process_docx_table(element) 
                signals.append({
                    'type': 'table',
                    'node': element, # Node preserved for Docx.Table reconstruction
                    'text': table_html
                })
                
        return signals

    def _xml_flag_enabled(self, element, tag_name: str) -> bool:
        """Interpret Word on/off flags like <w:b/> and <w:b w:val=\"0\"/> safely."""
        flag = element.find(qn(tag_name))
        if flag is None:
            return False

        val = flag.get(qn('w:val'))
        if val is None:
            return True

        return str(val).strip().lower() not in {'0', 'false', 'off', 'no'}

    def _restart_numbering(self, paragraph):
        """Forcefully breaks the link to previous lists and starts fresh."""
        pPr = paragraph._element.get_or_add_pPr()
        
        # 1. COMPLETELY WIPE any existing numbering properties from the Style
        for e in pPr.findall(qn('w:numPr')):
            pPr.remove(e)

        # 2. Create a brand new property block
        new_numPr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0') # Level 0 (top level)
        
        # 3. Use an even wider range for the counter to ensure it's unique
        if not hasattr(self, '_num_id_counter'):
            self._num_id_counter = 1000 # Higher start point for "Clean Slate"
        self._num_id_counter += 1
        
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), str(self._num_id_counter))
        
        new_numPr.append(ilvl)
        new_numPr.append(numId)
        pPr.append(new_numPr)

    def _process_docx_table(self, table_elem):
        """Convers Word table XML to simplified HTML for chat preview."""
        table_html = "<table>"
        for row in table_elem.findall(qn('w:tr')):
            table_html += "<tr>"
            for cell in row.findall(qn('w:tc')):
                cell_text = ""
                for p in cell.findall(qn('w:p')):
                    for r in p.findall(qn('w:r')):
                        t_node = r.find(qn('w:t'))
                        if t_node is not None: cell_text += t_node.text
                    cell_text += " "
                table_html += f"<td>{cell_text.strip()}</td>"
            table_html += "</tr>"
        table_html += "</table>"
        return table_html


